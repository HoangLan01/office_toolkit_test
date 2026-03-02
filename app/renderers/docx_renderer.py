from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.agent.office_writer_agent import DocumentDraft

class DocxRenderer:
    def render(self, draft: DocumentDraft, output_path: str):
        """
        Renders the DocumentDraft into a DOCX file using python-docx.
        """
        doc = Document()
        
        # Add Title
        title_para = doc.add_heading(draft.title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for section in draft.sections:
            # Add Heading
            doc.add_heading(section.heading, level=section.level)
            
            # Add Paragraphs
            for para in section.paragraphs:
                p = doc.add_paragraph(para.text)
                if para.style and para.style.lower() in ["quote", "emphasis"]:
                    # Quick styling based on the hint
                    p.style = 'Intense Quote' if 'quote' in para.style.lower() else 'Normal'
                    
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))
        print(f"DOCX saved to {output_file}")
